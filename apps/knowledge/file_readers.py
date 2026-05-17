import io
import logging

logger = logging.getLogger(__name__)


def read_file(file_path: str, source_type: str) -> str:
    """
    Read a file from disk and extract all text.
    source_type: pdf | docx | txt | xlsx | doc
    """
    try:
        if source_type == "pdf":
            return _read_pdf(file_path)
        elif source_type in ("docx", "doc"):
            return _read_docx(file_path)
        elif source_type == "txt":
            return _read_txt(file_path)
        elif source_type == "xlsx":
            return _read_xlsx(file_path)
        else:
            logger.warning(f"Unsupported file type: {source_type}")
            return ""
    except Exception as e:
        logger.error(f"Error reading file {file_path}: {e}")
        raise


def read_file_from_bytes(content: bytes, mime_type: str) -> str:
    """
    Read file from raw bytes (used for Google Drive / OneDrive downloads).
    mime_type examples: application/pdf, application/vnd.openxmlformats...
    """
    try:
        if "pdf" in mime_type:
            return _read_pdf_bytes(content)
        elif "wordprocessingml" in mime_type or "msword" in mime_type:
            return _read_docx_bytes(content)
        elif "spreadsheetml" in mime_type or "excel" in mime_type:
            return _read_xlsx_bytes(content)
        elif "text/plain" in mime_type:
            return content.decode("utf-8", errors="ignore")
        else:
            # try plain text as fallback
            return content.decode("utf-8", errors="ignore")
    except Exception as e:
        logger.error(f"Error reading file from bytes (mime={mime_type}): {e}")
        raise


# ─────────────────────────────────────────────
# PDF
# ─────────────────────────────────────────────

def _read_pdf(file_path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(file_path)
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages.append(f"[Page {i + 1}]\n{text.strip()}")
    result = "\n\n".join(pages)
    logger.info(f"PDF: extracted {len(result)} chars from {len(reader.pages)} pages")
    return result


def _read_pdf_bytes(content: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(content))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages.append(f"[Page {i + 1}]\n{text.strip()}")
    return "\n\n".join(pages)


# ─────────────────────────────────────────────
# DOCX
# ─────────────────────────────────────────────

def _read_docx(file_path: str) -> str:
    from docx import Document
    doc = Document(file_path)
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # also read tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    result = "\n".join(parts)
    logger.info(f"DOCX: extracted {len(result)} chars")
    return result


def _read_docx_bytes(content: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(content))
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


# ─────────────────────────────────────────────
# TXT
# ─────────────────────────────────────────────

def _read_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read().strip()
    logger.info(f"TXT: extracted {len(content)} chars")
    return content


# ─────────────────────────────────────────────
# XLSX
# ─────────────────────────────────────────────

def _read_xlsx(file_path: str) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"[Sheet: {sheet_name}]")
        for row in ws.iter_rows(values_only=True):
            row_values = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if row_values:
                parts.append(" | ".join(row_values))

    result = "\n".join(parts)
    logger.info(f"XLSX: extracted {len(result)} chars from {len(wb.sheetnames)} sheets")
    return result


def _read_xlsx_bytes(content: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"[Sheet: {sheet_name}]")
        for row in ws.iter_rows(values_only=True):
            row_values = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if row_values:
                parts.append(" | ".join(row_values))
    return "\n".join(parts)
